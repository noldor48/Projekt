import sqlite3
import sqlite3
from tkinter import *
from tkinter import ttk
from tkinter import messagebox
from docx import Document
from docx.shared import Inches

#Datubazes atvershana
conn = sqlite3.connect('Persons.db')
cursor = conn.cursor()

#Galvena funkcija
def cv_izveidoshana():
    def izveidoshana():
        id = id_entry.get()

        cursor.execute("SELECT * FROM Personas WHERE person_id LIKE ?", (f"%{id}%",))
        rezultati = cursor.fetchall()
        if rezultati:
            for r in rezultati:
                person = f"{r[1]} {r[2]}\n"
                info = f"Dzimsanas datums: {r[3]}\nVecums: {r[4]}\nDzimums: {r[5]}\nemail: {r[6]}\n"
                print(f"TYPE:{type(r[7])}")
                pic =r[7]
        cursor.execute("SELECT * FROM Sasniegumi WHERE person_id LIKE ?", (f"%{id}%",))
        sasniegumi = cursor.fetchall()
        if sasniegumi:
            sas = ""
            num=1
            for s in sasniegumi:
                sas += f"{str(num)}.Sasniegums\nDatums: {s[1]}\nVieta: {s[2]}\nSasniegums: {s[3]}\n"
                num+=1

        document = Document()
        document.add_heading(person, 0)
        document.add_picture(pic, width=Inches(1.25))
        document.add_paragraph(info)
        document.add_paragraph(sas)
        document.add_page_break()
        document.save('demo.docx')
        messagebox.showinfo("Veiksmīgi", "CV pievienots!")
        root.destroy()

    # Loga veidoshana
    root = Tk()
    root.title("Cv izveidoshana")
    root.iconbitmap(default="useri.ico")
    root.geometry(f"300x150+{int((root.winfo_screenwidth())/2)-150}+{int((root.winfo_screenheight())/2)-75}")
    root.configure(bg='#000000') 

    #Datus ievadishanas pieprasijums
    ttk.Label(root, text="id:").pack()
    id_entry = ttk.Entry(root)
    id_entry.pack()

    saglabat_btn = ttk.Button(root, text="Izveidot cv", command=izveidoshana)
    saglabat_btn.pack(pady=10)